import os
import io
import json
import logging
from typing import Dict, Any, List, Optional, Union, Tuple
from datetime import datetime
import mimetypes
import uuid

# Initialize logger early for use in import error handling
logger = logging.getLogger(__name__)

# File processing imports with optional fallbacks
try:
    import pandas as pd
except ImportError:
    pandas = None

try:
    from PIL import Image, ImageEnhance
except ImportError:
    Image = None
    ImageEnhance = None

try:
    import PyPDF2
except ImportError:
    PyPDF2 = None

try:
    import pdfplumber
except ImportError:
    pdfplumber = None

try:
    from docx import Document
except ImportError:
    Document = None

try:
    import openpyxl
except ImportError:
    openpyxl = None

import csv

# OCR and text processing (optional)
try:
    import pytesseract
except ImportError:
    pytesseract = None

try:
    import cv2
    import numpy as np
except ImportError:
    cv2 = None
    np = None

# Audio/Video processing (optional)
try:
    import speech_recognition as sr
except ImportError:
    sr = None

try:
    from moviepy.editor import VideoFileClip
except ImportError:
    VideoFileClip = None

try:
    import whisper
except (ImportError, TypeError, OSError, AttributeError) as e:
    # whisper may fail on Windows due to libc loading issues
    logger.warning(f"Failed to import whisper: {e}. Audio transcription will be unavailable.")
    whisper = None

# Utility imports
from pathlib import Path

from models.file_analysis import (
    ContentInsights, DocumentMetadata, ImageMetadata, 
    DataFileMetadata, AudioVideoMetadata
)

class FileParserRegistry:
    """Registry for different file format parsers"""
    
    def __init__(self):
        self.parsers = {}
        self._register_default_parsers()
    
    def register_parser(self, file_types: List[str], parser_class):
        """Register a parser for specific file types"""
        for file_type in file_types:
            self.parsers[file_type.lower()] = parser_class
    
    def _register_default_parsers(self):
        """Register default parsers for common file types"""
        self.register_parser(['pdf'], PDFParser)
        self.register_parser(['csv'], CSVParser)
        self.register_parser(['xlsx', 'xls'], ExcelParser)
        self.register_parser(['docx', 'doc'], WordDocumentParser)
        self.register_parser(['jpg', 'jpeg', 'png', 'gif', 'bmp', 'tiff'], ImageParser)
        self.register_parser(['txt', 'md', 'log'], TextParser)
        self.register_parser(['json'], JSONParser)
        self.register_parser(['mp3', 'wav', 'flac'], AudioParser)
        self.register_parser(['mp4', 'avi', 'mov', 'mkv'], VideoParser)
    
    def get_parser(self, file_type: str):
        """Get appropriate parser for file type"""
        return self.parsers.get(file_type.lower())

class BaseParser:
    """Base class for all file parsers"""
    
    def __init__(self, file_path: str):
        self.file_path = file_path
        self.file_type = Path(file_path).suffix[1:].lower()
    
    async def parse(self) -> Dict[str, Any]:
        """Parse file and return structured content"""
        raise NotImplementedError
    
    async def extract_metadata(self) -> Dict[str, Any]:
        """Extract file-specific metadata"""
        raise NotImplementedError
    
    def _get_file_stats(self) -> Dict[str, Any]:
        """Get basic file statistics"""
        stat = os.stat(self.file_path)
        return {
            "file_size": stat.st_size,
            "created_at": datetime.fromtimestamp(stat.st_ctime),
            "modified_at": datetime.fromtimestamp(stat.st_mtime),
            "file_type": self.file_type
        }

class PDFParser(BaseParser):
    """Parser for PDF documents"""
    
    async def parse(self) -> Dict[str, Any]:
        try:
            text_content = ""
            metadata = {}
            
            # Check if PDF libraries are available
            if not pdfplumber and not PyPDF2:
                text_content = f"PDF parsing not available - install PyPDF2 or pdfplumber"
                metadata = {"error": "PDF libraries not installed", "extraction_method": "none"}
                return {
                    "content": text_content,
                    "metadata": metadata,
                    "content_type": "document",
                    "word_count": 0,
                    "character_count": len(text_content)
                }
            
            # Try pdfplumber first (better text extraction)
            if pdfplumber:
                try:
                    with pdfplumber.open(self.file_path) as pdf:
                        page_texts = []
                        for page in pdf.pages:
                            page_text = page.extract_text()
                            if page_text:
                                page_texts.append(page_text)
                        
                        text_content = "\n\n".join(page_texts)
                        metadata.update({
                            "page_count": len(pdf.pages),
                            "extraction_method": "pdfplumber"
                        })
                except Exception as e:
                    logger.warning(f"pdfplumber failed for {self.file_path}: {e}")
                    text_content = ""
            
            # Fallback to PyPDF2 if pdfplumber failed or not available
            if not text_content and PyPDF2:
                try:
                    with open(self.file_path, 'rb') as file:
                        pdf_reader = PyPDF2.PdfReader(file)
                        page_texts = []
                        
                        for page in pdf_reader.pages:
                            page_text = page.extract_text()
                            if page_text:
                                page_texts.append(page_text)
                        
                        text_content = "\n\n".join(page_texts)
                        metadata.update({
                            "page_count": len(pdf_reader.pages),
                            "extraction_method": "pypdf2"
                        })
                except Exception as e:
                    logger.warning(f"PyPDF2 failed for {self.file_path}: {e}")
                    text_content = ""
            
            # OCR fallback if both text extraction methods failed
            if not text_content or len(text_content.strip()) < 50:
                logger.info(f"Attempting OCR fallback for {self.file_path}")
                try:
                    ocr_text = await self._extract_text_with_ocr()
                    if ocr_text and len(ocr_text.strip()) > 50:
                        text_content = ocr_text
                        metadata["extraction_method"] = "ocr_fallback"
                        metadata["ocr_used"] = True
                        logger.info(f"OCR successfully extracted {len(ocr_text)} characters")
                    else:
                        logger.warning("OCR extraction failed or returned minimal content")
                        if not text_content:
                            text_content = f"PDF text extraction failed - no readable text found"
                except Exception as ocr_error:
                    logger.warning(f"OCR fallback failed for {self.file_path}: {ocr_error}")
                    if not text_content:
                        text_content = f"PDF text extraction failed: {str(ocr_error)}"
            
            # Extract additional metadata
            doc_metadata = await self.extract_metadata()
            metadata.update(doc_metadata)
            
            return {
                "content": text_content,
                "metadata": metadata,
                "content_type": "document",
                "word_count": len(text_content.split()) if text_content else 0,
                "character_count": len(text_content)
            }
            
        except Exception as e:
            logger.error(f"Error parsing PDF {self.file_path}: {e}")
            raise
    
    async def _extract_text_with_ocr(self) -> str:
        """Extract text from PDF using OCR as fallback"""
        try:
            # Import OCR dependencies
            try:
                import fitz  # PyMuPDF
                import pytesseract
                from PIL import Image
                import io
            except ImportError as e:
                logger.warning(f"OCR dependencies not available: {e}")
                return ""
            
            # Open PDF with PyMuPDF
            doc = fitz.open(self.file_path)
            all_text = ""
            
            # Process first 10 pages to avoid excessive processing time
            max_pages = min(10, len(doc))
            
            for page_num in range(max_pages):
                try:
                    page = doc.load_page(page_num)
                    
                    # Convert page to image
                    pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))  # 2x scale for better OCR
                    img_data = pix.tobytes("png")
                    
                    # Use PIL to open the image
                    img = Image.open(io.BytesIO(img_data))
                    
                    # Use OCR to extract text
                    page_text = pytesseract.image_to_string(img, config='--oem 3 --psm 6')
                    
                    if page_text and page_text.strip():
                        all_text += f"\n--- Page {page_num + 1} ---\n"
                        all_text += page_text.strip()
                        
                except Exception as page_error:
                    logger.warning(f"OCR failed for page {page_num + 1}: {page_error}")
                    continue
            
            doc.close()
            return all_text.strip()
            
        except Exception as e:
            logger.error(f"OCR extraction failed: {e}")
            return ""
    
    async def extract_metadata(self) -> Dict[str, Any]:
        try:
            with open(self.file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                metadata = self._get_file_stats()
                
                # Extract PDF metadata if available
                if pdf_reader.metadata:
                    pdf_meta = pdf_reader.metadata
                    metadata.update({
                        "title": pdf_meta.get('/Title', ''),
                        "author": pdf_meta.get('/Author', ''),
                        "subject": pdf_meta.get('/Subject', ''),
                        "creator": pdf_meta.get('/Creator', ''),
                        "producer": pdf_meta.get('/Producer', ''),
                        "creation_date": pdf_meta.get('/CreationDate', ''),
                        "modification_date": pdf_meta.get('/ModDate', '')
                    })
                
                return metadata
                
        except Exception as e:
            logger.error(f"Error extracting PDF metadata: {e}")
            return self._get_file_stats()

class ImageParser(BaseParser):
    """Parser for image files with OCR capability"""
    
    async def parse(self) -> Dict[str, Any]:
        try:
            # Open and analyze image
            with Image.open(self.file_path) as img:
                # Basic image info
                metadata = {
                    "width": img.width,
                    "height": img.height,
                    "format": img.format,
                    "mode": img.mode
                }
                
                # Perform OCR to extract text
                text_content = ""
                try:
                    # Enhance image for better OCR
                    enhanced_img = ImageEnhance.Contrast(img).enhance(2.0)
                    enhanced_img = ImageEnhance.Sharpness(enhanced_img).enhance(2.0)
                    
                    # Extract text using OCR
                    text_content = pytesseract.image_to_string(enhanced_img)
                    metadata["has_text"] = bool(text_content.strip())
                    metadata["ocr_confidence"] = self._get_ocr_confidence(enhanced_img)
                    
                except Exception as ocr_error:
                    logger.warning(f"OCR failed for {self.file_path}: {ocr_error}")
                    metadata["has_text"] = False
                    metadata["ocr_error"] = str(ocr_error)
                
                # Get additional metadata
                img_metadata = await self.extract_metadata()
                metadata.update(img_metadata)
                
                return {
                    "content": text_content,
                    "metadata": metadata,
                    "content_type": "image",
                    "word_count": len(text_content.split()) if text_content else 0,
                    "character_count": len(text_content)
                }
                
        except Exception as e:
            logger.error(f"Error parsing image {self.file_path}: {e}")
            raise
    
    def _get_ocr_confidence(self, img) -> float:
        """Get OCR confidence score"""
        try:
            data = pytesseract.image_to_data(img, output_type=pytesseract.Output.DICT)
            confidences = [int(conf) for conf in data['conf'] if int(conf) > 0]
            return sum(confidences) / len(confidences) if confidences else 0.0
        except:
            return 0.0
    
    async def extract_metadata(self) -> Dict[str, Any]:
        metadata = self._get_file_stats()
        
        try:
            with Image.open(self.file_path) as img:
                # Extract EXIF data if available
                exif_data = {}
                if hasattr(img, '_getexif') and img._getexif():
                    exif = img._getexif()
                    if exif:
                        exif_data = {str(k): str(v) for k, v in exif.items()}
                
                metadata.update({
                    "dimensions": f"{img.width}x{img.height}",
                    "color_mode": img.mode,
                    "exif_data": exif_data
                })
                
        except Exception as e:
            logger.error(f"Error extracting image metadata: {e}")
        
        return metadata

class CSVParser(BaseParser):
    """Parser for CSV files"""
    
    async def parse(self) -> Dict[str, Any]:
        try:
            # Read CSV file
            df = pd.read_csv(self.file_path)
            
            # Generate content description
            content_description = self._describe_dataframe(df)
            
            # Extract metadata
            metadata = await self.extract_metadata()
            metadata.update({
                "row_count": len(df),
                "column_count": len(df.columns),
                "columns": df.columns.tolist(),
                "data_types": df.dtypes.astype(str).to_dict(),
                "missing_values": df.isnull().sum().to_dict(),
                "sample_data": df.head(3).to_dict('records') if len(df) > 0 else []
            })
            
            return {
                "content": content_description,
                "metadata": metadata,
                "content_type": "structured_data",
                "word_count": len(content_description.split()),
                "character_count": len(content_description)
            }
            
        except Exception as e:
            logger.error(f"Error parsing CSV {self.file_path}: {e}")
            raise
    
    def _describe_dataframe(self, df: pd.DataFrame) -> str:
        """Generate human-readable description of DataFrame"""
        description_parts = []
        
        description_parts.append(f"This CSV file contains {len(df)} rows and {len(df.columns)} columns.")
        
        if len(df.columns) > 0:
            description_parts.append(f"Columns: {', '.join(df.columns.tolist())}")
        
        # Describe data types
        numeric_cols = df.select_dtypes(include=[np.number]).columns.tolist()
        text_cols = df.select_dtypes(include=['object']).columns.tolist()
        
        if numeric_cols:
            description_parts.append(f"Numeric columns: {', '.join(numeric_cols)}")
        if text_cols:
            description_parts.append(f"Text columns: {', '.join(text_cols)}")
        
        # Basic statistics for numeric columns
        if numeric_cols and len(df) > 0:
            description_parts.append("\nNumeric column summaries:")
            for col in numeric_cols[:3]:  # Limit to first 3 numeric columns
                stats = df[col].describe()
                description_parts.append(
                    f"- {col}: mean={stats['mean']:.2f}, min={stats['min']:.2f}, max={stats['max']:.2f}"
                )
        
        return "\n".join(description_parts)
    
    async def extract_metadata(self) -> Dict[str, Any]:
        return self._get_file_stats()

class ExcelParser(BaseParser):
    """Parser for Excel files"""
    
    async def parse(self) -> Dict[str, Any]:
        try:
            # Read Excel file
            excel_file = pd.ExcelFile(self.file_path)
            sheet_names = excel_file.sheet_names
            
            content_parts = []
            all_metadata = {}
            total_rows = 0
            total_cols = 0
            
            for sheet_name in sheet_names:
                df = pd.read_excel(self.file_path, sheet_name=sheet_name)
                
                sheet_description = f"\nSheet '{sheet_name}':\n"
                sheet_description += self._describe_dataframe(df)
                content_parts.append(sheet_description)
                
                total_rows += len(df)
                total_cols = max(total_cols, len(df.columns))
                
                all_metadata[f"sheet_{sheet_name}"] = {
                    "row_count": len(df),
                    "column_count": len(df.columns),
                    "columns": df.columns.tolist()
                }
            
            content_description = f"Excel file with {len(sheet_names)} sheet(s):" + "".join(content_parts)
            
            metadata = await self.extract_metadata()
            metadata.update({
                "sheet_count": len(sheet_names),
                "sheet_names": sheet_names,
                "total_rows": total_rows,
                "total_columns": total_cols,
                "sheets_data": all_metadata
            })
            
            return {
                "content": content_description,
                "metadata": metadata,
                "content_type": "structured_data",
                "word_count": len(content_description.split()),
                "character_count": len(content_description)
            }
            
        except Exception as e:
            logger.error(f"Error parsing Excel {self.file_path}: {e}")
            raise
    
    def _describe_dataframe(self, df: pd.DataFrame) -> str:
        """Generate human-readable description of DataFrame"""
        description_parts = []
        
        description_parts.append(f"Contains {len(df)} rows and {len(df.columns)} columns.")
        
        if len(df.columns) > 0:
            description_parts.append(f"Columns: {', '.join(df.columns.tolist())}")
        
        return "\n".join(description_parts)
    
    async def extract_metadata(self) -> Dict[str, Any]:
        return self._get_file_stats()

class WordDocumentParser(BaseParser):
    """Parser for Word documents"""
    
    async def parse(self) -> Dict[str, Any]:
        try:
            doc = Document(self.file_path)
            
            # Extract text from paragraphs
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            text_content = "\n\n".join(paragraphs)
            
            # Extract text from tables if any
            table_content = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells]
                    table_content.append(" | ".join(row_text))
            
            if table_content:
                text_content += "\n\nTables:\n" + "\n".join(table_content)
            
            metadata = await self.extract_metadata()
            metadata.update({
                "paragraph_count": len(paragraphs),
                "table_count": len(doc.tables),
                "has_tables": len(doc.tables) > 0
            })
            
            return {
                "content": text_content,
                "metadata": metadata,
                "content_type": "document",
                "word_count": len(text_content.split()) if text_content else 0,
                "character_count": len(text_content)
            }
            
        except Exception as e:
            logger.error(f"Error parsing Word document {self.file_path}: {e}")
            raise
    
    async def extract_metadata(self) -> Dict[str, Any]:
        metadata = self._get_file_stats()
        
        try:
            doc = Document(self.file_path)
            
            # Extract document properties if available
            core_props = doc.core_properties
            metadata.update({
                "title": core_props.title or "",
                "author": core_props.author or "",
                "subject": core_props.subject or "",
                "created": core_props.created,
                "modified": core_props.modified,
                "category": core_props.category or "",
                "comments": core_props.comments or ""
            })
            
        except Exception as e:
            logger.error(f"Error extracting Word document metadata: {e}")
        
        return metadata

class TextParser(BaseParser):
    """Parser for plain text files"""
    
    async def parse(self) -> Dict[str, Any]:
        try:
            with open(self.file_path, 'r', encoding='utf-8', errors='ignore') as file:
                content = file.read()
            
            metadata = await self.extract_metadata()
            metadata.update({
                "encoding": "utf-8",
                "line_count": content.count('\n') + 1
            })
            
            return {
                "content": content,
                "metadata": metadata,
                "content_type": "text",
                "word_count": len(content.split()) if content else 0,
                "character_count": len(content)
            }
            
        except Exception as e:
            logger.error(f"Error parsing text file {self.file_path}: {e}")
            raise
    
    async def extract_metadata(self) -> Dict[str, Any]:
        return self._get_file_stats()

class JSONParser(BaseParser):
    """Parser for JSON files"""
    
    async def parse(self) -> Dict[str, Any]:
        try:
            with open(self.file_path, 'r', encoding='utf-8') as file:
                data = json.load(file)
            
            # Create human-readable description of JSON structure
            content_description = self._describe_json_structure(data)
            
            metadata = await self.extract_metadata()
            metadata.update({
                "json_type": type(data).__name__,
                "top_level_keys": list(data.keys()) if isinstance(data, dict) else [],
                "array_length": len(data) if isinstance(data, list) else None,
                "structure": self._analyze_json_structure(data)
            })
            
            return {
                "content": content_description,
                "metadata": metadata,
                "content_type": "structured_data",
                "word_count": len(content_description.split()),
                "character_count": len(content_description)
            }
            
        except Exception as e:
            logger.error(f"Error parsing JSON file {self.file_path}: {e}")
            raise
    
    def _describe_json_structure(self, data) -> str:
        """Generate human-readable description of JSON structure"""
        if isinstance(data, dict):
            description = f"JSON object with {len(data)} keys: {', '.join(list(data.keys())[:10])}"
            if len(data) > 10:
                description += "..."
        elif isinstance(data, list):
            description = f"JSON array with {len(data)} items"
            if len(data) > 0:
                item_types = set(type(item).__name__ for item in data[:5])
                description += f" (item types: {', '.join(item_types)})"
        else:
            description = f"JSON value of type {type(data).__name__}"
        
        return description
    
    def _analyze_json_structure(self, data, max_depth=3, current_depth=0) -> Dict[str, Any]:
        """Analyze JSON structure recursively"""
        if current_depth >= max_depth:
            return {"truncated": True}
        
        if isinstance(data, dict):
            return {
                "type": "object",
                "key_count": len(data),
                "keys": list(data.keys())[:10],
                "nested_structures": {
                    k: self._analyze_json_structure(v, max_depth, current_depth + 1)
                    for k, v in list(data.items())[:5]
                }
            }
        elif isinstance(data, list):
            return {
                "type": "array",
                "length": len(data),
                "item_types": list(set(type(item).__name__ for item in data[:10]))
            }
        else:
            return {"type": type(data).__name__, "value": str(data)[:100]}
    
    async def extract_metadata(self) -> Dict[str, Any]:
        return self._get_file_stats()

class AudioParser(BaseParser):
    """Parser for audio files with speech-to-text"""
    
    async def parse(self) -> Dict[str, Any]:
        try:
            # Initialize speech recognizer
            r = sr.Recognizer()
            
            # Try to load and transcribe audio
            text_content = ""
            metadata = await self.extract_metadata()
            
            try:
                # Load audio file
                with sr.AudioFile(self.file_path) as source:
                    audio = r.record(source)
                
                # Perform speech recognition
                text_content = r.recognize_google(audio)
                metadata.update({
                    "has_speech": True,
                    "transcription_method": "google",
                    "transcript_available": True
                })
                
            except sr.UnknownValueError:
                text_content = "No speech detected in audio file"
                metadata.update({"has_speech": False, "transcript_available": False})
            except sr.RequestError as e:
                logger.warning(f"Speech recognition service error: {e}")
                text_content = "Speech recognition service unavailable"
                metadata.update({"has_speech": "unknown", "transcript_available": False})
            except Exception as e:
                logger.warning(f"Audio transcription failed: {e}")
                text_content = "Audio transcription failed"
                metadata.update({"has_speech": "unknown", "transcript_available": False})
            
            return {
                "content": text_content,
                "metadata": metadata,
                "content_type": "audio",
                "word_count": len(text_content.split()) if text_content else 0,
                "character_count": len(text_content)
            }
            
        except Exception as e:
            logger.error(f"Error parsing audio file {self.file_path}: {e}")
            raise
    
    async def extract_metadata(self) -> Dict[str, Any]:
        metadata = self._get_file_stats()
        
        try:
            # Try to get audio properties using moviepy
            from moviepy.editor import AudioFileClip
            
            with AudioFileClip(self.file_path) as audio:
                metadata.update({
                    "duration": audio.duration,
                    "fps": getattr(audio, 'fps', None),
                    "audio_channels": getattr(audio, 'nchannels', None)
                })
                
        except Exception as e:
            logger.warning(f"Could not extract audio metadata: {e}")
        
        return metadata

class VideoParser(BaseParser):
    """Parser for video files with audio transcription"""
    
    async def parse(self) -> Dict[str, Any]:
        try:
            text_content = ""
            metadata = await self.extract_metadata()
            
            try:
                # Extract audio from video and transcribe
                with VideoFileClip(self.file_path) as video:
                    if video.audio:
                        # Save audio temporarily
                        temp_audio_path = f"/tmp/{uuid.uuid4()}.wav"
                        video.audio.write_audiofile(temp_audio_path, verbose=False, logger=None)
                        
                        # Transcribe audio
                        r = sr.Recognizer()
                        with sr.AudioFile(temp_audio_path) as source:
                            audio = r.record(source)
                        
                        text_content = r.recognize_google(audio)
                        
                        # Clean up temp file
                        os.remove(temp_audio_path)
                        
                        metadata.update({
                            "has_audio": True,
                            "transcript_available": True,
                            "transcription_method": "google"
                        })
                    else:
                        text_content = "No audio track found in video"
                        metadata.update({"has_audio": False, "transcript_available": False})
                        
            except Exception as e:
                logger.warning(f"Video transcription failed: {e}")
                text_content = "Video transcription failed"
                metadata.update({"has_audio": "unknown", "transcript_available": False})
            
            return {
                "content": text_content,
                "metadata": metadata,
                "content_type": "video",
                "word_count": len(text_content.split()) if text_content else 0,
                "character_count": len(text_content)
            }
            
        except Exception as e:
            logger.error(f"Error parsing video file {self.file_path}: {e}")
            raise
    
    async def extract_metadata(self) -> Dict[str, Any]:
        metadata = self._get_file_stats()
        
        try:
            with VideoFileClip(self.file_path) as video:
                metadata.update({
                    "duration": video.duration,
                    "fps": video.fps,
                    "width": video.w,
                    "height": video.h,
                    "has_audio": video.audio is not None
                })
                
        except Exception as e:
            logger.warning(f"Could not extract video metadata: {e}")
        
        return metadata

# Global parser registry instance
parser_registry = FileParserRegistry()

async def parse_file(file_path: str, file_type: str = None) -> Dict[str, Any]:
    """
    Parse a file and extract its content and metadata
    
    Args:
        file_path: Path to the file to parse
        file_type: Optional file type override
    
    Returns:
        Dictionary containing parsed content and metadata
    """
    if not file_type:
        file_type = Path(file_path).suffix[1:].lower()
    
    parser_class = parser_registry.get_parser(file_type)
    
    if not parser_class:
        # Fallback to text parser for unknown types
        logger.warning(f"No specific parser found for {file_type}, using text parser")
        parser_class = TextParser
    
    parser = parser_class(file_path)
    return await parser.parse()